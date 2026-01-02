import os
from django.shortcuts import render
from django.http import HttpResponse
from django.conf import settings
from googletrans import Translator
from docx import Document

translator = Translator()


def index(request):
    return render(request, "index.html")


def translate_file(request):
    if request.method != "POST":
        return HttpResponse("Invalid request method")

    uploaded_file = request.FILES.get("file")
    src_lang = request.POST.get("source_lang")
    dest_lang = request.POST.get("dest_lang")

    if not uploaded_file:
        return HttpResponse("No file uploaded")

    file_name, file_ext = os.path.splitext(uploaded_file.name)

    # ===================== TXT FILE =====================
    if file_ext.lower() == ".txt":
        content = uploaded_file.read().decode("utf-8")

        translated_text = translator.translate(
            content, src=src_lang, dest=dest_lang
        ).text

        response = HttpResponse(translated_text, content_type="text/plain")
        response["Content-Disposition"] = (
            f'attachment; filename="{file_name}_{dest_lang}.txt"'
        )
        return response

    # ===================== DOCX FILE =====================
    elif file_ext.lower() == ".docx":
        doc = Document(uploaded_file)

        # Translate paragraph runs (keeps images & formatting)
        for para in doc.paragraphs:
            for run in para.runs:
                if run.text.strip():
                    translated = translator.translate(
                        run.text, src=src_lang, dest=dest_lang
                    ).text
                    run.text = translated

        # Translate table text (very important)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if run.text.strip():
                                translated = translator.translate(
                                    run.text, src=src_lang, dest=dest_lang
                                ).text
                                run.text = translated

        output_path = os.path.join(
            settings.BASE_DIR, f"{file_name}_{dest_lang}.docx"
        )
        doc.save(output_path)

        with open(output_path, "rb") as f:
            response = HttpResponse(
                f.read(),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            response["Content-Disposition"] = (
                f'attachment; filename="{file_name}_{dest_lang}.docx"'
            )
            return response

    # ===================== UNSUPPORTED =====================
    else:
        return HttpResponse("Unsupported file type")
